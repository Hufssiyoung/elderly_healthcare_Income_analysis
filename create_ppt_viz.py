"""
PPT_Visualization.docx 생성 스크립트
PPT에 들어갈 핵심 시각화를 모아서 Word 문서로 저장합니다.

실행 방법:
    python create_ppt_viz.py
"""
import os, io, warnings
warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import seaborn as sns
from scipy.stats import linregress
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler, label_binarize
from sklearn.linear_model import Ridge, Lasso, LogisticRegression
from sklearn.svm import SVC
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (
    mean_squared_error, r2_score, f1_score, accuracy_score,
    roc_auc_score, confusion_matrix, roc_curve, auc, classification_report,
)
from xgboost import XGBRegressor, XGBClassifier
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RANDOM_STATE = 42
DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
OUT_PATH = os.path.join(os.path.dirname(__file__), "PPT_Visualization.docx")

FEAT_COLS = [
    "고령화율", "고령화율_변화", "독거노인비율",
    "비율_75-79세", "비율_80-84세", "비율_85세이상",
    "인프라갭", "포화도", "시설수", "수요공급갭지수",
]
TARGET = "요양수요가속도"
CLASS_ORDER = ["저위험", "중위험", "고위험"]
class_map = {c: i for i, c in enumerate(CLASS_ORDER)}
inv_map = {i: c for c, i in class_map.items()}
CLUSTER_NAMES = {
    0: "초고령·인프라포화", 1: "고령화가속·인프라갭",
    2: "고령화안정·인프라충족", 3: "저고령·수요태동", 4: "고령독거·수요폭증",
}
K_OPTIMAL = 5
PALETTE_C = ["blue", "green", "yellow", "orange", "red"]


# ── 폰트 설정 ──────────────────────────────────────────────────────────────
def setup_font():
    for fp in [
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
    ]:
        if os.path.exists(fp):
            fm.fontManager.addfont(fp)
            plt.rcParams["font.family"] = fm.FontProperties(fname=fp).get_name()
            break
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 150


# ── 그림 → BytesIO 저장 헬퍼 ──────────────────────────────────────────────
def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf


# ── Docx 헬퍼 ─────────────────────────────────────────────────────────────
def add_section_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_figure(doc, buf, caption, width=6.0):
    doc.add_picture(buf, width=Inches(width))
    p = doc.add_paragraph(f"▲ {caption}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


# ── 데이터 로드 ──────────────────────────────────────────────────────────
def load_data():
    df = pd.read_csv(os.path.join(DATA_DIR, "df_analysis.csv"), encoding="utf-8-sig")
    dc = pd.read_csv(os.path.join(DATA_DIR, "df_criteria.csv"), encoding="utf-8-sig")
    df_b = df.merge(dc[["행정동코드", "위험등급", "IMD_score"]], on="행정동코드", how="inner")
    df_b["위험등급"] = pd.Categorical(df_b["위험등급"], categories=CLASS_ORDER, ordered=True)
    return df, df_b


def build_pipe(model, scale=True):
    steps = [("scaler", StandardScaler())] if scale else []
    steps.append(("model", model))
    return Pipeline(steps)


def rmse(yt, yp):
    return np.sqrt(mean_squared_error(yt, yp))


# ════════════════════════════════════════════════════════════════════════════
# 섹션 1. EDA
# ════════════════════════════════════════════════════════════════════════════
def section_eda(doc, df, df_b):
    print("  [EDA] 시각화 생성 중...")
    add_section_heading(doc, "1. 탐색적 데이터 분석 (EDA)", level=1)

    # 1-1. 상관계수 히트맵
    corr_cols = FEAT_COLS + [TARGET]
    corr = df[corr_cols].corr()
    fig, ax = plt.subplots(figsize=(11, 9))
    mask = np.triu(np.ones_like(corr, dtype=bool), k=1)
    sns.heatmap(corr, mask=mask, annot=True, fmt=".2f", cmap="RdYlBu_r",
                center=0, vmin=-1, vmax=1, square=True, linewidths=0.5,
                annot_kws={"size": 8}, ax=ax)
    ax.set_title("변수 간 Pearson 상관계수 히트맵", fontsize=13, fontweight="bold")
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "X·Y 변수 간 Pearson 상관계수 히트맵 (하삼각)")

    # 1-2. 상위/하위 10개 행정동
    df_y = df[["행정동명", "자치구명", TARGET]].sort_values(TARGET, ascending=False)
    fig, axes = plt.subplots(1, 2, figsize=(15, 5))
    top10 = df_y.head(10)
    axes[0].barh(top10["행정동명"] + " (" + top10["자치구명"] + ")", top10[TARGET], color="tomato")
    axes[0].set_title("요양수요가속도 상위 10개 행정동\n(독거노인 급증 지역)")
    axes[0].invert_yaxis()
    axes[0].set_xlabel(TARGET)
    bot10 = df_y.tail(10)
    axes[1].barh(bot10["행정동명"] + " (" + bot10["자치구명"] + ")", bot10[TARGET], color="steelblue")
    axes[1].set_title("요양수요가속도 하위 10개 행정동\n(독거노인 감소/정체 지역)")
    axes[1].invert_yaxis()
    axes[1].set_xlabel(TARGET)
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "요양수요가속도 상위/하위 10개 행정동")

    # 1-3. IMD 위험등급 파이차트
    d_labels = ["D1 노인장기요양\n수급자비율(30%)", "D2 장기요양시설부족(25%)",
                "D3 미충족의료율(25%)", "D4 기초수급비율(20%)"]
    d_weights = [0.30, 0.25, 0.25, 0.20]
    fig, ax = plt.subplots(figsize=(7, 5))
    ax.pie(d_weights, labels=d_labels, autopct="%1.0f%%", startangle=90,
            textprops={"fontsize": 10})
    ax.set_title("IMD 구성 영역별 가중치", fontsize=12, fontweight="bold")
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "IMD(고령지역불평등지수) 구성 지표별 가중치", width=4.5)


# ════════════════════════════════════════════════════════════════════════════
# 섹션 2. Track A — 회귀
# ════════════════════════════════════════════════════════════════════════════
def section_track_a(doc, df):
    print("  [Track A] 모델 학습 및 시각화 생성 중...")
    add_section_heading(doc, "2. Track A — 요양수요가속도 회귀분석", level=1)

    X = df[FEAT_COLS].copy()
    y = df[TARGET].copy()
    y_bins = pd.qcut(y, q=5, labels=False, duplicates="drop")
    X_tr, X_te, y_tr, y_te = train_test_split(
        X, y, test_size=0.2, random_state=RANDOM_STATE, stratify=y_bins)
    for p in [X_tr, X_te, y_tr, y_te]:
        p.reset_index(drop=True, inplace=True)

    default_models = {
        "Ridge": build_pipe(Ridge(alpha=1.0), scale=True),
        "Lasso": build_pipe(Lasso(alpha=0.1, max_iter=10000), scale=True),
        "DecisionTree": build_pipe(
            __import__("sklearn.tree", fromlist=["DecisionTreeRegressor"]).DecisionTreeRegressor(random_state=RANDOM_STATE), scale=False),
        "RandomForest": build_pipe(
            __import__("sklearn.ensemble", fromlist=["RandomForestRegressor"]).RandomForestRegressor(n_estimators=100, random_state=RANDOM_STATE, n_jobs=-1), scale=False),
        "XGBoost": build_pipe(
            XGBRegressor(n_estimators=100, random_state=RANDOM_STATE, verbosity=0, n_jobs=-1), scale=False),
    }

    base_results = []
    for name, pipe in default_models.items():
        from sklearn.base import clone
        m = clone(pipe)
        m.fit(X_tr, y_tr)
        p_tr = m.predict(X_tr)
        p_te = m.predict(X_te)
        base_results.append({
            "모델": name,
            "Train_RMSE": rmse(y_tr, p_tr), "Test_RMSE": rmse(y_te, p_te),
            "Train_R2": r2_score(y_tr, p_tr), "Test_R2": r2_score(y_te, p_te),
            "Overfit_Gap": rmse(y_te, p_te) - rmse(y_tr, p_tr),
        })
    df_base = pd.DataFrame(base_results).sort_values("Test_R2", ascending=False)

    # 2-1. 기본 모델 비교
    sorted_names = df_base["모델"].tolist()
    x = np.arange(len(sorted_names))
    w = 0.38
    fig, axes = plt.subplots(1, 3, figsize=(17, 5))
    for ax, (tr_c, te_c, title) in zip(axes, [
        ("Train_RMSE", "Test_RMSE", "Test RMSE"),
        ("Train_R2", "Test_R2", "Test R²"),
        ("Overfit_Gap", None, "과적합 격차 (Test - Train RMSE)"),
    ]):
        if te_c is None:
            vals = [df_base[df_base["모델"] == n][tr_c].values[0] for n in sorted_names]
            colors = ["tomato" if v > 3 else "seagreen" for v in vals]
            ax.bar(sorted_names, vals, color=colors, alpha=0.85, edgecolor="white")
            ax.axhline(0, color="gray", lw=1, ls="--")
        else:
            tr_v = [df_base[df_base["모델"] == n][tr_c].values[0] for n in sorted_names]
            te_v = [df_base[df_base["모델"] == n][te_c].values[0] for n in sorted_names]
            ax.bar(x - w/2, tr_v, w, label="Train", color="steelblue", alpha=0.85, edgecolor="white")
            ax.bar(x + w/2, te_v, w, label="Test", color="tomato", alpha=0.85, edgecolor="white")
            ax.set_xticks(x)
            ax.legend(fontsize=9)
            if "R²" in title:
                ax.axhline(0, color="orange", lw=1, ls="--")
        ax.set_xticklabels(sorted_names, rotation=20, ha="right")
        ax.set_title(title, fontsize=11, fontweight="bold")
    plt.suptitle("기본 모델 비교 (기본 하이퍼파라미터, Train/Test)", fontsize=13, fontweight="bold", y=1.02)
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track A 기본 모델 성능 비교 (RMSE / R² / 과적합 격차)")

    # 2-2. 튜닝 모델 훈련 (best params)
    tuned_models = {
        "Ridge": build_pipe(Ridge(alpha=10), scale=True),
        "Lasso": build_pipe(Lasso(alpha=0.1, max_iter=10000), scale=True),
        "DecisionTree": build_pipe(
            __import__("sklearn.tree", fromlist=["DecisionTreeRegressor"]).DecisionTreeRegressor(
                max_depth=5, min_samples_leaf=5, min_samples_split=20, random_state=RANDOM_STATE), scale=False),
        "RandomForest": build_pipe(
            __import__("sklearn.ensemble", fromlist=["RandomForestRegressor"]).RandomForestRegressor(
                n_estimators=200, max_depth=None, max_features="log2",
                min_samples_leaf=2, min_samples_split=2, random_state=RANDOM_STATE, n_jobs=-1), scale=False),
        "XGBoost": build_pipe(
            XGBRegressor(n_estimators=300, max_depth=3, learning_rate=0.1,
                          subsample=0.8, colsample_bytree=1.0, reg_alpha=0.1,
                          random_state=RANDOM_STATE, verbosity=0, n_jobs=-1), scale=False),
    }
    tuned, best_name, best_model = {}, None, None
    best_rmse = 1e9
    for name, pipe in tuned_models.items():
        from sklearn.base import clone
        m = clone(pipe)
        m.fit(X_tr, y_tr)
        tuned[name] = m
        te_r = rmse(y_te, m.predict(X_te))
        if te_r < best_rmse:
            best_rmse = te_r
            best_name = name
            best_model = m

    # 2-3. 실제 vs 예측 (best model)
    y_pred = best_model.predict(X_te)
    residuals = y_te.values - y_pred
    mn, mx = min(y_te.min(), y_pred.min()), max(y_te.max(), y_pred.max())
    fig, ax = plt.subplots(figsize=(7, 6))
    sc = ax.scatter(y_te, y_pred, alpha=0.6, s=30,
                     c=np.abs(residuals), cmap="RdYlGn_r", edgecolors="none")
    ax.plot([mn, mx], [mn, mx], "k--", lw=1.5, label="완벽한 예측")
    plt.colorbar(sc, ax=ax, shrink=0.8, label="|잔차|")
    ax.set_xlabel("실제 요양수요가속도")
    ax.set_ylabel("예측 요양수요가속도")
    r2 = r2_score(y_te, y_pred)
    ax.set_title(f"{best_name} — 실제 vs 예측  (Test R²={r2:.3f})", fontsize=12, fontweight="bold")
    ax.legend()
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), f"Track A 최적 모델({best_name}) 실제 vs 예측 산점도", width=4.5)

    # 2-4. SHAP 전역 중요도
    try:
        import shap
        if hasattr(best_model, "named_steps"):
            core = best_model.named_steps[list(best_model.named_steps.keys())[-1]]
            if "scaler" in best_model.named_steps:
                X_te_t = pd.DataFrame(best_model.named_steps["scaler"].transform(X_te), columns=FEAT_COLS)
            else:
                X_te_t = X_te.copy()
        else:
            core, X_te_t = best_model, X_te.copy()
        exp = shap.TreeExplainer(core)
        sv = exp.shap_values(X_te_t)
        mean_abs = np.abs(sv).mean(axis=0)
        shap_df = pd.DataFrame({"피처": FEAT_COLS, "Mean |SHAP|": mean_abs}).sort_values("Mean |SHAP|")
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.barh(shap_df["피처"], shap_df["Mean |SHAP|"],
                 color=plt.cm.RdBu_r(np.linspace(0.2, 0.8, len(FEAT_COLS))),
                 edgecolor="white", alpha=0.85)
        ax.set_title(f"{best_name} — SHAP 전역 피처 중요도", fontsize=12, fontweight="bold")
        ax.set_xlabel("Mean |SHAP value|")
        plt.tight_layout()
        add_figure(doc, fig_to_bytes(fig), f"Track A SHAP 전역 피처 중요도 ({best_name})")
    except Exception as e:
        print(f"    SHAP 오류: {e}")


# ════════════════════════════════════════════════════════════════════════════
# 섹션 3. Track B — 분류
# ════════════════════════════════════════════════════════════════════════════
def section_track_b(doc, df_b):
    print("  [Track B] 모델 학습 및 시각화 생성 중...")
    add_section_heading(doc, "3. Track B — 위험등급 분류분석", level=1)

    X = df_b[FEAT_COLS].copy()
    y = df_b["위험등급"].copy()
    X_tr, X_te, y_tr, y_te = train_test_split(
        X, y, test_size=0.2, random_state=RANDOM_STATE, stratify=y)
    for p in [X_tr, X_te, y_tr, y_te]:
        p.reset_index(drop=True, inplace=True)
    y_tr_enc = y_tr.map(class_map)
    y_te_enc = y_te.map(class_map)
    sorted_labels = sorted(CLASS_ORDER)
    xgb_col_order = [class_map[c] for c in sorted_labels]

    # 3-1. 피처 vs 위험등급 상관계수
    y_num = y.cat.codes
    corr = df_b[FEAT_COLS].corrwith(y_num).sort_values(key=abs, ascending=False)
    fig, ax = plt.subplots(figsize=(10, 4))
    bar_colors = ["tomato" if v > 0 else "steelblue" for v in corr.values]
    ax.barh(corr.index, corr.values, color=bar_colors, edgecolor="white", alpha=0.85)
    ax.axvline(0, color="gray", lw=1)
    ax.set_title("피처 vs 위험등급 상관계수 (저위험=0, 중위험=1, 고위험=2)",
                  fontsize=12, fontweight="bold")
    ax.set_xlabel("Pearson r")
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track B 피처 vs 위험등급 Pearson 상관계수")

    # 3-2. 모델 훈련 (best params)
    tuned = {
        "LogisticRegression": build_pipe(
            LogisticRegression(C=1.0, max_iter=2000, random_state=RANDOM_STATE), scale=True),
        "SVM_RBF": build_pipe(
            SVC(kernel="rbf", C=50, gamma="scale", probability=True, random_state=RANDOM_STATE), scale=True),
        "RandomForest": build_pipe(
            RandomForestClassifier(n_estimators=300, max_depth=7, min_samples_leaf=5,
                                    max_features=0.5, random_state=RANDOM_STATE, n_jobs=-1), scale=False),
        "XGBoost": build_pipe(
            XGBClassifier(n_estimators=200, max_depth=3, learning_rate=0.05,
                           subsample=1.0, colsample_bytree=0.8, reg_alpha=0.1, reg_lambda=2.0,
                           random_state=RANDOM_STATE, verbosity=0, n_jobs=-1, eval_metric="mlogloss"), scale=False),
    }
    test_results, test_preds, test_probs = [], {}, {}
    for name, pipe in tuned.items():
        from sklearn.base import clone
        m = clone(pipe)
        y_fit = y_tr_enc if name == "XGBoost" else y_tr
        m.fit(X_tr, y_fit)
        p_te = m.predict(X_te)
        prob_te = m.predict_proba(X_te)
        if name == "XGBoost":
            p_te = pd.Categorical([inv_map[i] for i in p_te], categories=CLASS_ORDER, ordered=True)
            prob_te = prob_te[:, xgb_col_order]
        test_preds[name] = p_te
        test_probs[name] = prob_te
        test_results.append({
            "모델": name,
            "Train_F1": f1_score(y_tr, pd.Categorical([inv_map[i] for i in m.predict(X_tr)],
                                  categories=CLASS_ORDER, ordered=True) if name == "XGBoost"
                        else m.predict(X_tr), average="macro"),
            "Test_F1": f1_score(y_te, p_te, average="macro"),
            "Test_Acc": accuracy_score(y_te, p_te),
            "Test_AUC": roc_auc_score(y_te, prob_te, multi_class="ovr", average="macro"),
        })
    df_test = pd.DataFrame(test_results).sort_values("Test_F1", ascending=False)
    best_name = df_test.iloc[0]["모델"]

    # 3-3. 기본 모델 비교 (F1/Acc/과적합)
    model_names = list(tuned.keys())
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    x = np.arange(len(model_names))
    w = 0.35
    for ax, (tr_c, te_c, title) in zip(axes, [
        ("Train_F1", "Test_F1", "Macro F1 (Train vs Test)"),
        ("Test_AUC", None, "Test ROC-AUC (OvR)"),
    ]):
        if te_c is None:
            vals = [df_test[df_test["모델"] == n][tr_c].values[0] for n in model_names]
            ax.bar(model_names, vals, color="steelblue", alpha=0.85, edgecolor="white")
        else:
            tr_v = [df_test[df_test["모델"] == n][tr_c].values[0] for n in model_names]
            te_v = [df_test[df_test["모델"] == n][te_c].values[0] for n in model_names]
            ax.bar(x - w/2, tr_v, w, label="Train", color="steelblue", alpha=0.85, edgecolor="white")
            ax.bar(x + w/2, te_v, w, label="Test", color="tomato", alpha=0.85, edgecolor="white")
            ax.set_xticks(x)
            ax.legend()
            ax.set_ylim(0, 1.1)
        ax.set_xticklabels(model_names, rotation=15, ha="right")
        ax.set_title(title, fontsize=11, fontweight="bold")
    plt.suptitle("Track B 기본 모델 성능 비교", fontsize=13, fontweight="bold", y=1.02)
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track B 모델 Macro F1 / ROC-AUC 비교")

    # 3-4. 최적모델 Confusion Matrix
    cm = confusion_matrix(y_te, test_preds[best_name], labels=CLASS_ORDER)
    cm_norm = cm.astype(float) / cm.sum(axis=1, keepdims=True)
    fig, ax = plt.subplots(figsize=(6, 5))
    sns.heatmap(cm_norm, annot=cm, fmt="d", cmap="Blues",
                xticklabels=CLASS_ORDER, yticklabels=CLASS_ORDER,
                ax=ax, linewidths=0.5, vmin=0, vmax=1,
                annot_kws={"size": 14, "weight": "bold"})
    f1_val = df_test[df_test["모델"] == best_name]["Test_F1"].values[0]
    ax.set_title(f"{best_name} — Confusion Matrix\nMacro F1={f1_val:.3f}",
                  fontsize=12, fontweight="bold")
    ax.set_xlabel("예측 클래스")
    ax.set_ylabel("실제 클래스")
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), f"Track B 최적 모델({best_name}) 혼동 행렬", width=4.0)

    # 3-5. 최종 성능 비교 막대
    metrics_p = ["Test_F1", "Test_Acc", "Test_AUC"]
    labels_p = ["Macro F1", "Accuracy", "ROC-AUC"]
    x2 = np.arange(len(metrics_p))
    w2 = 0.2
    palette = ["steelblue", "seagreen", "goldenrod", "tomato"]
    fig, ax = plt.subplots(figsize=(10, 4))
    for i, name in enumerate(model_names):
        vals = [df_test[df_test["모델"] == name][m].values[0] for m in metrics_p]
        ax.bar(x2 + (i - 1.5) * w2, vals, w2,
                label=name, color=palette[i], alpha=0.85, edgecolor="white")
    ax.set_xticks(x2)
    ax.set_xticklabels(labels_p)
    ax.set_ylim(0, 1.15)
    ax.legend(fontsize=9)
    ax.axhline(1/3, color="gray", lw=1, ls="--", alpha=0.5)
    ax.set_title("Track B 최종 모델 성능 비교 (튜닝 후 Test Set)", fontsize=12, fontweight="bold")
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track B 전체 모델 Macro F1 / Accuracy / ROC-AUC 비교")


# ════════════════════════════════════════════════════════════════════════════
# 섹션 4. Track C — 군집화
# ════════════════════════════════════════════════════════════════════════════
def section_track_c(doc, df):
    print("  [Track C] 군집화 및 시각화 생성 중...")
    add_section_heading(doc, "4. Track C — 지역 군집분석 (K-Means K=5)", level=1)

    X_raw = df[FEAT_COLS].copy()
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X_raw)
    km = KMeans(n_clusters=K_OPTIMAL, n_init=50, random_state=RANDOM_STATE)
    df["cluster_km"] = km.fit_predict(X_scaled)
    df["cluster_label"] = df["cluster_km"].map(CLUSTER_NAMES)
    acc_means = df.groupby("cluster_km")["요양수요가속도"].mean()
    sorted_clusters = acc_means.sort_values().index.tolist()
    cluster_sizes = df["cluster_km"].value_counts()
    pca_model = PCA(n_components=2, random_state=RANDOM_STATE)
    X_pca = pca_model.fit_transform(X_scaled)

    # 4-1. 레이더 차트
    cluster_means_raw = df.groupby("cluster_km")[FEAT_COLS].mean()
    cmin, cmax = cluster_means_raw.min(), cluster_means_raw.max()
    cluster_means_norm = (cluster_means_raw - cmin) / (cmax - cmin + 1e-9)
    N = len(FEAT_COLS)
    angles = np.linspace(0, 2*np.pi, N, endpoint=False).tolist()
    angles += angles[:1]
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw={"polar": True})
    for i, cid in enumerate(sorted_clusters):
        color = PALETTE_C[i % len(PALETTE_C)]
        values = cluster_means_norm.loc[cid, FEAT_COLS].values.tolist()
        values += values[:1]
        n_dong = cluster_sizes.get(cid, 0)
        ax.plot(angles, values, "o-", lw=2.5, color=color,
                 label=f"C{cid}: {CLUSTER_NAMES[cid]} (n={n_dong})")
        ax.fill(angles, values, alpha=0.12, color=color)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(FEAT_COLS, fontsize=10)
    ax.set_ylim(0, 1)
    ax.set_title("클러스터별 취약성 프로파일\n(정규화 [0,1])", fontsize=13, fontweight="bold", pad=25)
    ax.legend(loc="upper right", bbox_to_anchor=(1.45, 1.15), fontsize=9)
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track C 클러스터별 취약성 레이더 차트 (정규화)")

    # 4-2. 클러스터 히트맵 Z-score
    cluster_means_z = cluster_means_raw.apply(lambda col: (col - col.mean()) / (col.std() + 1e-9))
    cluster_means_z.index = [f"C{cid}\n{CLUSTER_NAMES[cid]}" for cid in cluster_means_z.index]
    fig, ax = plt.subplots(figsize=(11, 6))
    sns.heatmap(cluster_means_z, annot=True, fmt=".2f", cmap="RdBu_r", center=0,
                linewidths=1.5, ax=ax, annot_kws={"size": 11, "weight": "bold"},
                cbar_kws={"label": "Z-score"}, vmin=-2, vmax=2)
    ax.set_title("클러스터별 피처 Z-score 프로파일\n빨간색=상대 높음, 파란색=상대 낮음",
                  fontsize=12, fontweight="bold")
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track C 클러스터 프로파일 히트맵 (Z-score)")

    # 4-3. PCA 2D
    var_ratio = pca_model.explained_variance_ratio_
    fig, ax = plt.subplots(figsize=(9, 6))
    for i, cid in enumerate(sorted_clusters):
        mask = df["cluster_km"] == cid
        ax.scatter(X_pca[mask, 0], X_pca[mask, 1],
                    c=PALETTE_C[i % len(PALETTE_C)], s=35, alpha=0.75, edgecolors="none",
                    label=f"C{cid}: {CLUSTER_NAMES[cid]} (n={mask.sum()})")
    centers_pca = pca_model.transform(km.cluster_centers_)
    for cid in range(K_OPTIMAL):
        rank_i = sorted_clusters.index(cid)
        ax.scatter(centers_pca[cid, 0], centers_pca[cid, 1],
                    c=PALETTE_C[rank_i % len(PALETTE_C)], s=250, marker="*",
                    edgecolors="white", linewidths=1.5, zorder=5)
        ax.annotate(f"C{cid}", (centers_pca[cid, 0], centers_pca[cid, 1]),
                     fontsize=10, fontweight="bold", xytext=(5, 5), textcoords="offset points")
    ax.set_xlabel(f"PC1 ({var_ratio[0]*100:.1f}%)", fontsize=11)
    ax.set_ylabel(f"PC2 ({var_ratio[1]*100:.1f}%)", fontsize=11)
    ax.set_title(f"PCA 2D 군집 시각화 (K=5, ★=센트로이드)\n"
                  f"PC1+PC2={sum(var_ratio)*100:.1f}% 분산 설명", fontsize=12, fontweight="bold")
    ax.legend(fontsize=9)
    ax.grid(alpha=0.2)
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track C PCA 2D 군집 시각화")

    # 4-4. 클러스터별 핵심 지표 2×2
    metrics_c = [
        ("고령화율", "(%)", "고령화율 평균"),
        ("독거노인비율", "(%)", "독거노인비율 평균"),
        ("인프라갭", "", "인프라갭 평균"),
        ("요양수요가속도", "", "요양수요가속도 평균 (Y변수)"),
    ]
    x_pos = np.arange(K_OPTIMAL)
    bar_colors = [PALETTE_C[sorted_clusters.index(cid) % len(PALETTE_C)] for cid in range(K_OPTIMAL)]
    xlabels = [f"C{cid}\n{CLUSTER_NAMES[cid]}" for cid in range(K_OPTIMAL)]
    fig, axes = plt.subplots(2, 2, figsize=(15, 10))
    for ax, (feat, unit, title) in zip(axes.flatten(), metrics_c):
        means = [df[df["cluster_km"] == cid][feat].mean() for cid in range(K_OPTIMAL)]
        stds  = [df[df["cluster_km"] == cid][feat].std()  for cid in range(K_OPTIMAL)]
        ax.bar(x_pos, means, color=bar_colors, alpha=0.85, edgecolor="white", width=0.6,
                yerr=stds, capsize=5, error_kw={"ecolor": "#94a3b8", "lw": 1.5})
        overall = df[feat].mean()
        ax.axhline(overall, color="#fbbf24", ls="--", lw=1.8, label=f"전체 평균 {overall:.2f}")
        for i, (v, s) in enumerate(zip(means, stds)):
            ax.text(i, v + s*1.1 + abs(max(means))*0.01, f"{v:.2f}",
                     ha="center", fontsize=9, fontweight="bold")
        ax.set_xticks(x_pos)
        ax.set_xticklabels(xlabels, fontsize=8)
        ax.set_title(title, fontsize=11, fontweight="bold")
        ax.legend(fontsize=9)
        ax.grid(axis="y", alpha=0.2)
    plt.suptitle("클러스터별 핵심 지표 종합 비교\n(막대=평균, 에러바=표준편차, 점선=서울 전체 평균)",
                  fontsize=14, fontweight="bold", y=1.01)
    plt.tight_layout()
    add_figure(doc, fig_to_bytes(fig), "Track C 클러스터별 핵심 지표 2×2 비교 차트")


# ════════════════════════════════════════════════════════════════════════════
# 메인
# ════════════════════════════════════════════════════════════════════════════
def main():
    print("PPT_Visualization.docx 생성 시작...")
    setup_font()

    df, df_b = load_data()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"

    # 표지
    title = doc.add_heading("고령 취약계층 헬스케어 수요 예측 플랫폼", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("PPT 발표 자료 — 핵심 시각화 모음")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph("서울시 420개 행정동 | 2023년 기준 | EDA · Track A(회귀) · Track B(분류) · Track C(군집화)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    section_eda(doc, df, df_b)
    doc.add_page_break()
    section_track_a(doc, df)
    doc.add_page_break()
    section_track_b(doc, df_b)
    doc.add_page_break()
    section_track_c(doc, df.copy())

    doc.save(OUT_PATH)
    print(f"\n✅ 저장 완료: {OUT_PATH}")
    print("총 섹션: EDA(3) + Track A(4) + Track B(5) + Track C(4) = 16개 시각화")


if __name__ == "__main__":
    main()
